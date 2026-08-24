"""
Document Preprocessor Module for Royalty Statement Parser
===========================================================
Extracts text, tables, structural content, and image renderings from diverse file formats:
- PDF (text, tables, or scanned pages)
- CSV / TSV
- XLS / XLSX
- DOC / DOCX
- TXT
- Images (PNG, JPG, JPEG)
"""
import io
import os
import re
import csv
import base64
from typing import Dict, Any, List, Optional, Tuple


def detect_file_type(filename: str, content_bytes: bytes) -> str:
    """
    Detect document format from filename extension and magic bytes.
    Returns one of: 'pdf', 'csv', 'tsv', 'xlsx', 'xls', 'docx', 'txt', 'image'.
    """
    ext = os.path.splitext(filename)[1].lower().strip(".")

    # Magic byte checks
    if content_bytes.startswith(b"%PDF"):
        return "pdf"
    if content_bytes.startswith(b"\x89PNG") or content_bytes.startswith(b"\xff\xd8\xff"):
        return "image"
    if content_bytes.startswith(b"PK\x03\x04"):
        if ext in ("xlsx", "xlsm"):
            return "xlsx"
        if ext == "docx":
            return "docx"
        return "xlsx" if "sheet" in filename.lower() else "docx"

    if ext in ("csv", "tsv", "txt", "xlsx", "xls", "docx", "pdf", "png", "jpg", "jpeg"):
        if ext in ("png", "jpg", "jpeg"):
            return "image"
        return ext

    # Fallback text/binary check
    try:
        content_bytes.decode("utf-8")
        return "tsv" if "\t" in filename.lower() else "csv"
    except UnicodeDecodeError:
        return "unknown"


def extract_content_from_file(filename: str, content_bytes: bytes) -> Dict[str, Any]:
    """
    Extract structured text, tables, and image payload from raw file bytes.

    Returns:
        {
            "file_type": str,
            "filename": str,
            "text_content": str,          # Structured text representation
            "table_data": List[List[str]], # Structured table rows if available
            "images": List[Dict[str, str]],# List of base64 data URIs for vision LLMs
            "page_count": int,
            "has_visual_tables": bool
        }
    """
    file_type = detect_file_type(filename, content_bytes)
    result: Dict[str, Any] = {
        "file_type": file_type,
        "filename": filename,
        "text_content": "",
        "table_data": [],
        "images": [],
        "page_count": 1,
        "has_visual_tables": False
    }

    if file_type in ("csv", "tsv", "txt"):
        text = _extract_from_text_or_csv(content_bytes, file_type)
        result["text_content"] = text

    elif file_type == "xlsx":
        text, tables = _extract_from_excel(content_bytes)
        result["text_content"] = text
        result["table_data"] = tables

    elif file_type == "docx":
        text = _extract_from_docx(content_bytes)
        result["text_content"] = text

    elif file_type == "pdf":
        text, tables, images, pages = _extract_from_pdf(content_bytes)
        result["text_content"] = text
        result["table_data"] = tables
        result["images"] = images
        result["page_count"] = pages
        if images and not text.strip():
            result["has_visual_tables"] = True

    elif file_type == "image":
        b64 = base64.b64encode(content_bytes).decode("utf-8")
        mime = "image/png" if filename.lower().endswith(".png") else "image/jpeg"
        result["images"] = [{"data_uri": f"data:{mime};base64,{b64}", "page": 1}]
        result["has_visual_tables"] = True
        result["text_content"] = f"[Scanned Image Document: {filename}]"

    else:
        # Fallback text decoding
        try:
            text = content_bytes.decode("utf-8", errors="replace")
            result["text_content"] = text
        except Exception:
            result["text_content"] = f"[Binary File: {filename}]"

    return result


def _extract_from_text_or_csv(content_bytes: bytes, file_type: str) -> str:
    """Decode text, CSV, or TSV file bytes."""
    for encoding in ("utf-8", "utf-8-sig", "latin1", "cp1252"):
        try:
            return content_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content_bytes.decode("utf-8", errors="replace")


def _extract_from_excel(content_bytes: bytes) -> Tuple[str, List[List[str]]]:
    """Extract sheets and tables from XLSX using openpyxl."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content_bytes), data_only=True)
        lines = []
        all_rows = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"--- Sheet: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                if not any(row):
                    continue
                row_str = [str(cell) if cell is not None else "" for cell in row]
                lines.append(" | ".join(row_str))
                all_rows.append(row_str)

        return "\n".join(lines), all_rows
    except Exception as e:
        print(f"[Preprocessor] Excel extraction fallback notice: {e}")
        text = content_bytes.decode("utf-8", errors="replace")
        return text, []


def _extract_from_docx(content_bytes: bytes) -> str:
    """Extract text and tables from DOCX using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(content_bytes))
        lines = []

        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())

        for table in doc.tables:
            lines.append("--- Table ---")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))

        return "\n".join(lines)
    except Exception as e:
        print(f"[Preprocessor] DOCX extraction fallback notice: {e}")
        return content_bytes.decode("utf-8", errors="replace")


def _extract_from_pdf(content_bytes: bytes) -> Tuple[str, List[List[str]], List[Dict[str, str]], int]:
    """
    Extract native text, tables, and render scanned pages from PDF bytes.
    Prefers pdfplumber / pypdf, with visual image fallback if pages are scanned.
    """
    text_lines = []
    tables = []
    images = []
    page_count = 1

    # 1. Try pdfplumber for table & text extraction
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
            page_count = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages, 1):
                p_text = page.extract_text()
                if p_text and p_text.strip():
                    text_lines.append(f"--- Page {page_num} ---")
                    text_lines.append(p_text.strip())

                p_tables = page.extract_tables()
                for tbl in p_tables:
                    for r in tbl:
                        if r and any(r):
                            cleaned_row = [str(c).strip() if c else "" for c in r]
                            tables.append(cleaned_row)

        full_text = "\n".join(text_lines)
        if full_text.strip():
            return full_text, tables, images, page_count
    except Exception as e:
        print(f"[Preprocessor] pdfplumber notice: {e}")

    # 2. Try pypdf fallback
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content_bytes))
        page_count = len(reader.pages)
        for page_num, page in enumerate(reader.pages, 1):
            p_text = page.extract_text()
            if p_text and p_text.strip():
                text_lines.append(f"--- Page {page_num} ---")
                text_lines.append(p_text.strip())

        full_text = "\n".join(text_lines)
        if full_text.strip():
            return full_text, tables, images, page_count
    except Exception as e:
        print(f"[Preprocessor] pypdf notice: {e}")

    # 3. If no native text extracted (scanned PDF), attempt image conversion using Pillow/pypdfium2 if available
    try:
        import pypdfium2
        pdf_file = pypdfium2.PdfDocument(content_bytes)
        page_count = len(pdf_file)
        for i in range(min(page_count, 5)):
            image = pdf_file[i].render(scale=2).to_pil()
            buf = io.BytesIO()
            image.save(buf, format="JPEG")
            b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
            images.append({"data_uri": f"data:image/jpeg;base64,{b64}", "page": i + 1})
    except Exception as e:
        print(f"[Preprocessor] PDF image render notice: {e}")

    return "\n".join(text_lines), tables, images, page_count


def sample_content_for_llm(content_str: str, max_chars: int = 15000) -> str:
    """
    Intelligently sample large documents to send header, first ~60 rows,
    last ~15 summary rows, and totals to the LLM.
    """
    lines = [l for l in content_str.splitlines() if l.strip()]
    if not lines or len(content_str) <= max_chars:
        return content_str

    header_lines = lines[:65]
    tail_lines = lines[-20:]

    sample = "\n".join(header_lines) + f"\n\n... [Truncated {len(lines) - 85} intermediate rows] ...\n\n" + "\n".join(tail_lines)
    return sample[:max_chars]
